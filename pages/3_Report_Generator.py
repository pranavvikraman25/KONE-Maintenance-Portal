# app.py
import streamlit as st
import pandas as pd
import json
from io import BytesIO
from datetime import datetime, date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from backend.report_utils import save_report
import os


st.set_page_config(page_title="KPI Word Report — EQ & Date Filters", layout="wide")
st.title("KPI Report Generator — Filter by EQ & Date Range")

# ------------------ KPI MAP & Colors ------------------
KPI_MAP = {
    "Door Friction": "doorFriction",
    "Door Speed Error": "cumulativeDoorSpeedError",
    "Landing Door Lock Hook Closing Time": "lockHookClosingTime",
    "Landing Door Lock Hook Open Time": "lockHookTime",
    "Maximum Force During Coupler Compress": "maximumForceDuringCompress",
    "Landing Door Lock Roller Clearance": "landingDoorLockRollerClearance"
}

# These are hex strings without '#'
COLOR_HEADER = "D9D9D9"
COLOR_GOOD = "92D050"   # green
COLOR_ACTION = "FFF200" # yellow
COLOR_BORDER = "000000"

THRESHOLDS = {
    "doorFriction": (30.0, 50.0),
    "cumulativeDoorSpeedError": (0.05, 0.08),
    "lockHookClosingTime": (0.2, 0.6),
    "lockHookTime": (0.2, 0.6),
    "maximumForceDuringCompress": (5.0, 28.0),
    "landingDoorLockRollerClearance": (0.0, 0.029)
}

# ------------------ Helpers ------------------
def safe_float(v):
    try:
        if pd.isna(v):
            return None
        return float(v)
    except Exception:
        return None

def set_cell_shading_hex(cell, hex_color):
    """Shade a docx table cell with a hex color (without #)."""
    if not hex_color:
        return
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="000000", size="6"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ("top","left","bottom","right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        tcPr.append(tag)

def detect_column(cols, candidates):
    """Return first column name that contains any candidate substring (case-insensitive)."""
    lower = {c.lower(): c for c in cols}
    for cand in candidates:
        for k in lower:
            if cand in k:
                return lower[k]
    return None

def read_uploaded_file(uploaded):
    """Robust read for xlsx/xls/csv/json"""
    name = uploaded.name.lower()
    try:
        if name.endswith(".xlsx"):
            # openpyxl is preferred for modern files
            df = pd.read_excel(uploaded, engine="openpyxl")
        elif name.endswith(".xls"):
            # xlrd required for old xls
            df = pd.read_excel(uploaded, engine="xlrd")
        elif name.endswith(".csv"):
            df = pd.read_csv(uploaded)
        elif name.endswith(".json"):
            # load JSON: handle dict-of-lists or list-of-dicts
            data = json.load(uploaded)
            if isinstance(data, dict):
                # try to find a list-of-dicts inside
                for v in data.values():
                    if isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict):
                        df = pd.DataFrame(v)
                        break
                else:
                    df = pd.DataFrame([data])
            elif isinstance(data, list):
                df = pd.DataFrame(data)
            else:
                df = pd.DataFrame([data])
        else:
            # try pandas read_excel as fallback
            df = pd.read_excel(uploaded)
        # Normalize column names: strip
        df.columns = [c.strip() if isinstance(c, str) else c for c in df.columns]
        return df
    except Exception as e:
        raise

# ------------------ Upload UI ------------------
uploaded = st.file_uploader("Upload KPI Excel/CSV/JSON file", type=["xlsx","xls","csv","json"])

if not uploaded:
    st.info("Upload a data file (xlsx / xls / csv / json) that contains columns for EQ, Date, CKPI (or KPI name), Floor, and an Ave/Value column.")
    st.stop()

# read the file
try:
    df = read_uploaded_file(uploaded)
except Exception as e:
    st.error(f"Error reading the uploaded file: {e}")
    st.stop()

if df.empty:
    st.error("Uploaded file was read but contains no rows.")
    st.stop()

st.success("File read successfully.")

# ------------------ Detect important columns ------------------
cols = list(df.columns)

# heuristics for columns
eq_col = detect_column(cols, ["eq"])
date_col = detect_column(cols, ["date"])
ckpi_col = detect_column(cols, ["ckpi", "kpi"])
floor_col = detect_column(cols, ["floor"])
ave_col = detect_column(cols, ["ave", "value"])

# show detected columns to user (encourage verification)
st.write("Detected columns (automatically):")
st.write({
    "EQ column": eq_col,
    "Date column": date_col,
    "CKPI/KPI column": ckpi_col,
    "Floor column": floor_col,
    "Value/Average column": ave_col
})

# ------------------ Confirm or override column selection ------------------
st.subheader("Confirm or override column selection (if detection missed any)")

# Define defaults
default_eq = "eq"
default_date = "ckpi_statistics_date"
default_ckpi = "ckpi"

# Automatically preselect defaults if they exist in the columns
eq_col = st.selectbox("EQ column", options=[None] + cols, 
                      index=(cols.index(default_eq) + 1) if default_eq in cols else (cols.index(eq_col) + 1 if eq_col in cols else 0))

date_col = st.selectbox("Date column", options=[None] + cols, 
                        index=(cols.index(default_date) + 1) if default_date in cols else (cols.index(date_col) + 1 if date_col in cols else 0))

ckpi_col = st.selectbox("CKPI/KPI column", options=[None] + cols, 
                        index=(cols.index(default_ckpi) + 1) if default_ckpi in cols else (cols.index(ckpi_col) + 1 if ckpi_col in cols else 0))

# For Floor and Value, keep manual flexibility
floor_col = st.selectbox("Floor column", options=[None] + cols, 
                         index=(cols.index(floor_col) + 1) if floor_col in cols else 0)

ave_col = st.selectbox("Value/Ave column", options=[None] + cols, 
                       index=(cols.index(ave_col) + 1) if ave_col in cols else 0)


# Validate
required = {"EQ": eq_col, "Date": date_col, "CKPI": ckpi_col, "Floor": floor_col, "Value": ave_col}
missing = [k for k,v in required.items() if not v]
if missing:
    st.error(f"Missing column selections for: {', '.join(missing)}. Please select appropriate columns.")
    st.stop()

# Convert date column to datetime
df[date_col] = pd.to_datetime(df[date_col], errors="coerce")
if df[date_col].isna().all():
    st.error("Selected Date column could not be parsed to datetimes. Please choose a different Date column.")
    st.stop()

# Normalize CKPI text and create a key mapped to KPI_MAP values
df["__ckpi_norm"] = df[ckpi_col].astype(str).str.strip().str.lower()
kpi_map_lower = {k.lower(): v for k,v in KPI_MAP.items()}

# Offer KPI choices to include in the report (show friendly names)
available_kpis = []
for k in sorted(KPI_MAP.keys()):
    if any(df["__ckpi_norm"].str.contains(k.lower())):
        available_kpis.append(k)
# Allow user to pick from all KPI_MAP keys (in case naming mismatches)
if not available_kpis:
    available_kpis = list(KPI_MAP.keys())

st.subheader("Select KPIs to include")
selected_kpis = st.multiselect("Choose KPI(s)", options=list(KPI_MAP.keys()), default=list(KPI_MAP.keys()))

# EQ selection (support multi-select)
st.subheader("Select EQ(s) to include")
eq_choices = df[eq_col].dropna().unique().tolist()
selected_eqs = st.multiselect("EQ(s)", options=eq_choices, default=[eq_choices[0]] if eq_choices else [])

if not selected_eqs:
    st.error("Select at least one EQ.")
    st.stop()

# Date range selection
min_dt = df[date_col].min().date()
max_dt = df[date_col].max().date()
st.subheader("Date Range")
start_date, end_date = st.date_input("Choose start and end date", [min_dt, max_dt])

# Additional: choose whether to aggregate multiple entries per (EQ, date, floor, KPI) by mean or first
agg_choice = st.radio("If multiple readings exist for same EQ/date/floor/KPI, aggregate by:", options=["mean","first"], index=0)

# Apply filters
mask_eq = df[eq_col].isin(selected_eqs)
mask_date = (df[date_col].dt.date >= start_date) & (df[date_col].dt.date <= end_date)
filtered = df[mask_eq & mask_date].copy()

if filtered.empty:
    st.warning("No data after applying EQ and date filters. Adjust filters.")
    st.stop()

st.write(f"Filtered rows: {len(filtered)}")
st.dataframe(filtered.head(200))

# Preprocess filtered data: map CKPI names to internal keys if possible
# Create a column "__kpi_key" which maps to our internal KPI key (like "doorFriction") if possible
def map_ckpi_to_key(text):
    t = str(text).strip().lower()
    # exact match with display names
    for display, key in KPI_MAP.items():
        if display.lower() == t:
            return key
    # contains match
    for display, key in KPI_MAP.items():
        if display.lower() in t or t in display.lower():
            return key
    # contains the key itself
    for key in KPI_MAP.values():
        if key.lower() in t:
            return key
    return None

filtered["__kpi_key"] = filtered[ckpi_col].apply(map_ckpi_to_key)

# If user selected subset of KPIs, restrict keys accordingly
selected_keys = [KPI_MAP[k] for k in selected_kpis]

# Keep only rows which map to selected keys (if mapping succeeded). If mapping failed for a row but user selected that KPI,
# we will still try to treat the original CKPI text if it contains the selected KPI display substring.
def row_matches_selected(row):
    if row["__kpi_key"] in selected_keys:
        return True
    # fallback: if any selected_kpi display name appears in raw text
    raw = str(row[ckpi_col]).lower()
    for kdisplay in selected_kpis:
        if kdisplay.lower() in raw:
            return True
    return False

filtered = filtered[filtered.apply(row_matches_selected, axis=1)]
if filtered.empty:
    st.warning("After KPI selection and mapping, no rows remain. Try broadening KPI selection.")
    st.stop()

# We'll aggregate multiple rows per (eq, date, floor, kpi_key) as requested
group_cols = [eq_col, date_col, floor_col, "__kpi_key"]
if agg_choice == "mean":
    agg_df = filtered.groupby(group_cols, dropna=False)[ave_col].mean().reset_index()
else:
    # first
    agg_df = filtered.groupby(group_cols, dropna=False)[ave_col].first().reset_index()

# For rows where __kpi_key is None but textual CKPI matches selected display names, try to assign keys again
def ensure_kpi_key(row):
    if pd.notna(row["__kpi_key"]):
        return row["__kpi_key"]
    raw = str(row[ckpi_col]).lower()
    for display, key in KPI_MAP.items():
        if display.lower() in raw and display in selected_kpis:
            return key
    return None

# Merge agg_df back with original CKPI text for better reporting if needed
# We'll re-attach ckpi text by picking first ckpi text for each group
ckpi_text = filtered.groupby(group_cols)[ckpi_col].first().reset_index().rename(columns={ckpi_col: "__ckpi_text"})
agg_df = agg_df.merge(ckpi_text, on=group_cols, how="left")
agg_df["__kpi_key"] = agg_df.apply(ensure_kpi_key, axis=1)

# Convert date to date-only for grouping pages
agg_df["__date_only"] = pd.to_datetime(agg_df[date_col]).dt.date

# ------------------ Word Report Generation ------------------
st.subheader("Generate Word Report")
st.markdown("Report will produce one page per **(date, EQ)** combination. Each page shows selected KPIs across floors.")

def build_report_doc(agg_df, selected_eqs, selected_keys, selected_kpis):
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(14)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    # Dates and EQs ordering
    dates = sorted(agg_df["__date_only"].dropna().unique())
    eq_order = selected_eqs

    for d in dates:
        date_df = agg_df[agg_df["__date_only"] == d]
        for eq in eq_order:
            page_df = date_df[date_df[eq_col] == eq]
            if page_df.empty:
                continue

            # Title
            p = doc.add_paragraph()
            run = p.add_run(f"KPI Report — {d.isoformat()} — EQ: {eq}")
            run.bold = True
            run.font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")

            # Floors
            # Sort floors numerically if possible, else lexicographically
            floors = sorted(
                page_df[floor_col].dropna().unique(),
                key=lambda x: (int(x) if str(x).isdigit() else float('inf'), str(x))
            )

            if not floors:
                doc.add_paragraph("No floor data for this EQ/date.")
                doc.add_page_break()
                continue

            # Table headers
            cols_header = ["CKPI", "No Corrective Action Required", "Corrective Action Required"] + [f"Floor {f}" for f in floors]
            table = doc.add_table(rows=1, cols=len(cols_header))
            table.style = "Table Grid"
            table.autofit = True

            # Header formatting
            hdr_cells = table.rows[0].cells
            font_size = 10 if len(floors) <= 12 else max(6, 10 - ((len(floors) - 12) // 2))
            for i, name in enumerate(cols_header):
                hdr_cells[i].text = name
                set_cell_shading_hex(hdr_cells[i], COLOR_HEADER)
                set_cell_border(hdr_cells[i], color=COLOR_BORDER, size="8")
                para = hdr_cells[i].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in para.runs:
                    r.bold = True
                    r.font.size = Pt(font_size)

            # Fill rows for each KPI
            for display_name in selected_kpis:
                key = KPI_MAP.get(display_name)
                row_cells = table.add_row().cells
                row_cells[0].text = display_name

                # ------------------- NEW: Column 2 & 3 with thresholds -------------------
                low_high = THRESHOLDS.get(key)
                if low_high:
                    low, high = low_high
                    row_cells[1].text = f"{low} - {high}"  # No Corrective Action Required (green)
                    set_cell_shading_hex(row_cells[1], COLOR_GOOD)
                    set_cell_border(row_cells[1], color=COLOR_BORDER, size="6")
                    row_cells[2].text = f"< {low} or > {high}"  # Corrective Action Required (yellow)
                    set_cell_shading_hex(row_cells[2], COLOR_ACTION)
                    set_cell_border(row_cells[2], color=COLOR_BORDER, size="6")
                else:
                    row_cells[1].text = "-"
                    row_cells[2].text = "-"
                    set_cell_border(row_cells[1], color=COLOR_BORDER, size="6")
                    set_cell_border(row_cells[2], color=COLOR_BORDER, size="6")

                # Center & style text
                for ci in [1,2]:
                    for para in row_cells[ci].paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.bold = True
                            run.font.size = Pt(10)

                # ------------------- Floor values -------------------
                for fi, f in enumerate(floors):
                    target_cell = row_cells[3 + fi]
                    match = page_df[(page_df[floor_col] == f) & (page_df["__kpi_key"] == key)]
                    if match.empty:
                        val = None
                    else:
                        val = safe_float(match[ave_col].values[0])

                    if val is None:
                        target_cell.text = "-"
                        set_cell_shading_hex(target_cell, COLOR_ACTION)
                    else:
                        target_cell.text = f"{val:.2f}"
                        # color based on threshold
                        if low_high and (low <= val <= high):
                            set_cell_shading_hex(target_cell, COLOR_GOOD)
                        else:
                            set_cell_shading_hex(target_cell, COLOR_ACTION)
                    set_cell_border(target_cell, color=COLOR_BORDER, size="6")
                    for para in target_cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.bold = True
                            run.font.size = Pt(10)

            doc.add_page_break()

    return doc


# Build and download on click
if st.button("Generate & Download Word Report"):
    with st.spinner("Building Word report..."):
        try:
            doc = build_report_doc(agg_df, selected_eqs, selected_keys, selected_kpis)
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)

            # Define file name and bytes
            file_name = f"KPI_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            file_bytes = buf.getvalue()

            # 🔥 Auto-save to backend/reports/Report_Generator/
            saved_path = save_report(
                file_bytes,
                module_name="Report_Generator",
                filter_label="KPI_Report",
                extension="docx"
            )

            # ✅ Show success and confirmation
            st.success(f"✅ Report generated and saved: {os.path.basename(saved_path)}")

            # Download button
            st.download_button(
                "⬇️ Download Word Report",
                data=file_bytes,
                file_name=os.path.basename(saved_path),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error(f"Failed to build report: {e}")


