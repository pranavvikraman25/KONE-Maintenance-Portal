import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

st.set_page_config(page_title="Maintenance Tracker", layout="wide")
st.title("🧰 Maintenance Tracker — Technician Action Center")

st.markdown("""
Upload your **Actionable Report (Excel/CSV)** from Trend Analysis to mark maintenance actions.  
Each row represents a CKPI reading — mark:
- ✅ **Checked** if the task is complete or verified  
- ❌ **Wrong / Review** if the task needs attention  

Behavior:
- Default floor selection = Floor 1 (if present), else first floor.
- Select All / Deselect All act on the current filters and persist immediately.
- Export to Word (.docx) preserves the current selections and colors.
""")

# ---------------- Session state initialization ----------------
for k in ("uploaded_file", "maint_df", "maint_table_state"):
    if k not in st.session_state:
        st.session_state[k] = None

# ---------------- File upload / reuse ----------------
uploaded = st.file_uploader("📂 Upload Actionable Report (Excel/CSV)", type=["xlsx", "csv"])
if uploaded:
    try:
        df = pd.read_excel(uploaded) if uploaded.name.endswith(".xlsx") else pd.read_csv(uploaded)
        # normalize ckpi text
        if "ckpi" in df.columns:
            df["ckpi"] = df["ckpi"].astype(str).str.lower()
        # create stable uid if missing
        if "__uid" not in df.columns:
            df = df.reset_index(drop=True)
            df["__uid"] = df.index.astype(int)
        # ensure checkbox columns exist
        for col in ("✅ checked", "❌ wrong / review"):
            if col not in df.columns:
                df[col] = False
        st.session_state.maint_df = df.copy()
        st.session_state.maint_table_state = df.copy()
        st.session_state.uploaded_file = uploaded
        st.success("✅ File loaded and cached.")
    except Exception as e:
        st.error(f"Error reading file: {e}")
        st.stop()
else:
    if st.session_state.maint_df is not None:
        df = st.session_state.maint_df.copy()
        st.info("Using previously uploaded file from this session.")
    else:
        st.info("Upload an actionable report (Excel/CSV) to begin.")
        st.stop()

# ---------------- KEY CKPIS ----------------
KEY_CKPIS = [
    "doorfriction",
    "cumulativedoorspeederror",
    "lockhookclosingtime",
    "lockhooktime",
    "maximumforceduringcompress",
    "landingdoorlockrollerclearance",
]

# ---------------- Sidebar filters ----------------
st.sidebar.header("🔍 Filters")

# Equipment filter
eqs = sorted(df["eq"].dropna().unique()) if "eq" in df.columns else []
sel_eqs = st.sidebar.multiselect("Select Equipment(s)", eqs, default=eqs)

# KPI filter limited to the 6 keys that exist in data
available_ckpis = [k for k in KEY_CKPIS if k in df.get("ckpi", pd.Series([])).unique()]
sel_ckpis = st.sidebar.multiselect("Select KPI(s)", available_ckpis, default=available_ckpis)

# Date range (constrained to data)
if "ckpi_statistics_date" not in df.columns:
    st.error("Missing 'ckpi_statistics_date' column in your file.")
    st.stop()
df["ckpi_statistics_date"] = pd.to_datetime(df["ckpi_statistics_date"], errors="coerce")
if df["ckpi_statistics_date"].isna().all():
    st.error("Could not parse any dates in 'ckpi_statistics_date'.")
    st.stop()
min_d = df["ckpi_statistics_date"].min().date()
max_d = df["ckpi_statistics_date"].max().date()
sel_range = st.sidebar.date_input("Select Date Range", [min_d, max_d], min_value=min_d, max_value=max_d)
start_d, end_d = sel_range

# Floor filter: default = "1" if present else first floor
sel_floors = []
if "floor" in df.columns:
    floors = sorted(df["floor"].dropna().unique(), key=lambda x: str(x))
    floors_str = [str(f) for f in floors]
    # default floor selection logic
    default_floor = "1" if "1" in floors_str else (floors_str[0] if floors_str else None)
    sel_floors = st.sidebar.multiselect("Select Floor(s)", floors_str, default=[default_floor] if default_floor else [])
else:
    sel_floors = []

# ---------------- Build filtered view (based on maint_table_state for checkbox persistence) ----------------
base = st.session_state.maint_table_state.copy() if st.session_state.maint_table_state is not None else st.session_state.maint_df.copy()

# Ensure checkbox columns exist in base
for col in ("✅ checked", "❌ wrong / review"):
    if col not in base.columns:
        base[col] = False
# Ensure __uid exists
if "__uid" not in base.columns:
    base = base.reset_index(drop=True)
    base["__uid"] = base.index.astype(int)

# Apply equipment/kpi/date filters first
mask = pd.Series(True, index=base.index)
if sel_eqs:
    mask &= base["eq"].isin(sel_eqs)
if sel_ckpis:
    mask &= base["ckpi"].isin(sel_ckpis)
mask &= base["ckpi_statistics_date"].dt.date.between(start_d, end_d)

# Apply floor filter (compare as string)
if sel_floors and "floor" in base.columns:
    mask &= base["floor"].astype(str).isin(sel_floors)

df_filtered = base[mask].copy()

if df_filtered.empty:
    st.warning("No rows match the selected filters.")
    st.stop()

# ---------------- Variability index (same as before) ----------------
def calc_var(values):
    v = pd.to_numeric(values, errors="coerce").dropna()
    if len(v) < 4:
        return 0
    diffs = np.diff(v)
    return np.sum(np.diff(np.sign(diffs)) != 0) / len(v) * 100

if "ave" in df_filtered.columns:
    var_df = (
        df_filtered.groupby(["eq", "ckpi"])["ave"]
        .apply(calc_var)
        .reset_index()
        .rename(columns={"ave": "variability_index"})
    )
    # merge back into df_filtered by eq+ckpi, then update base copy to keep persistence consistent
    df_filtered = df_filtered.merge(var_df, on=["eq", "ckpi"], how="left")
    df_filtered["Priority Flag"] = np.where(df_filtered["variability_index"] > 30, "⚠️ High Variability", "")

# ---------------- Control buttons (Select All/Deselect All operate on filtered rows) ----------------
st.markdown("### 🧾 Maintenance Task Table")
c1, c2 = st.columns([1,1])
with c1:
    if st.button("☑️ Select All"):
        uids = df_filtered["__uid"].tolist()
        st.session_state.maint_table_state.loc[st.session_state.maint_table_state["__uid"].isin(uids), "✅ checked"] = True
        st.session_state.maint_table_state.loc[st.session_state.maint_table_state["__uid"].isin(uids), "❌ wrong / review"] = False
        st.experimental_rerun()
with c2:
    if st.button("🚫 Deselect All"):
        uids = df_filtered["__uid"].tolist()
        st.session_state.maint_table_state.loc[st.session_state.maint_table_state["__uid"].isin(uids), "✅ checked"] = False
        st.session_state.maint_table_state.loc[st.session_state.maint_table_state["__uid"].isin(uids), "❌ wrong / review"] = False
        st.experimental_rerun()

# show info about selected floors (optional clarity)
if sel_floors:
    st.info(f"📍 Showing Floor(s): {', '.join(sel_floors)}")
else:
    st.info("📍 Showing all floors (no floor filter)")

# ---------------- Display table using data_editor but persist by __uid ----------------
# Prepare display frame (only filtered rows)
display_cols = list(df_filtered.columns)  # includes __uid and checkbox cols
display_df = df_filtered[display_cols].copy()

# Ensure checkbox columns exist
for cc in ("✅ checked", "❌ wrong / review"):
    if cc not in display_df.columns:
        display_df[cc] = False

# Use data_editor with checkbox column_config (fast)
edited = st.data_editor(
    display_df,
    use_container_width=True,
    hide_index=True,
    num_rows="dynamic",
    key="maint_table",
    column_config={
        "✅ checked": st.column_config.CheckboxColumn("✅ Checked", help="Mark as completed"),
        "❌ wrong / review": st.column_config.CheckboxColumn("❌ Review", help="Mark for recheck"),
    },
)

# Enforce mutual exclusivity and write changes back to session_state.maint_table_state using __uid
# 'edited' contains the filtered rows only; update maint_table_state by __uid
if not edited.empty:
    # ensure boolean fill
    edited["✅ checked"] = edited["✅ checked"].fillna(False).astype(bool)
    edited["❌ wrong / review"] = edited["❌ wrong / review"].fillna(False).astype(bool)

    for idx, row in edited.iterrows():
        uid = row["__uid"]
        # mutual exclusivity
        if row["✅ checked"] and row["❌ wrong / review"]:
            # prioritize the last clicked — but we can't detect last click here; prefer ✅ as previous behavior
            row["❌ wrong / review"] = False
        # update master state row by uid
        mask_uid = st.session_state.maint_table_state["__uid"] == uid
        for col in edited.columns:
            st.session_state.maint_table_state.loc[mask_uid, col] = row[col]

# Update display copy after persist (so highlight and export use latest)
final_display = st.session_state.maint_table_state.loc[st.session_state.maint_table_state["__uid"].isin(display_df["__uid"])].copy()

# ---------------- Highlight rows visually in Streamlit ----------------
def highlight_action(row):
    if row.get("✅ checked"):
        return ["background-color: #b5e7a0"] * len(row)
    if row.get("❌ wrong / review"):
        return ["background-color: #f4a6a6"] * len(row)
    return [""] * len(row)

st.dataframe(final_display.style.apply(highlight_action, axis=1), use_container_width=True)

# ---------------- Export to Word (Landscape), include all columns ----------------
if st.button("✅ Submit and Generate Word Report"):
    master = st.session_state.maint_table_state.copy()
    checked_df = master[master["✅ checked"] == True]
    wrong_df = master[master["❌ wrong / review"] == True]

    if checked_df.empty and wrong_df.empty:
        st.warning("No tasks selected. Please mark items before exporting.")
    else:
        doc = Document()
        # set landscape
        section = doc.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        doc.add_heading("Maintenance Review Report", level=1)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")

        # include all columns in master (ordered)
        headers = list(master.columns)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)

        # add rows in the same order as master (so export is stable)
        for _, r in master.iterrows():
            row_cells = table.add_row().cells
            for i, h in enumerate(headers):
                cell_text = r.get(h, "")
                row_cells[i].text = "" if pd.isna(cell_text) else str(cell_text)
                # shading per row
                if r.get("✅ checked"):
                    color = "C6EFCE"
                elif r.get("❌ wrong / review"):
                    color = "FFC7CE"
                else:
                    color = "FFFFFF"
                row_cells[i]._tc.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color))
                )

        # save to buffer and provide download
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.download_button(
            "💾 Download Maintenance Report (Word)",
            data=buf,
            file_name=f"Maintenance_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
